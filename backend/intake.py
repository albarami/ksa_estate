"""Document intake: .docx land opportunity → structured data → Geoportal parcel.

Pipeline: parse docx → Claude extraction → coordinate resolution → Geoportal merge.
"""

from __future__ import annotations

import io
import json
import logging
import re
from typing import Any

import httpx
from anthropic import AsyncAnthropic
from docx import Document
from pyproj import Transformer

from backend.data_fetch_http import fetch_land_object
from backend.geocode import find_parcel_at_coords, parse_coordinates

log = logging.getLogger("intake")

# Saudi cadastral projection: AIN-ABD / UTM zone 38N
_utm_to_wgs84 = Transformer.from_crs("EPSG:20438", "EPSG:4326", always_xy=True)

EXTRACTION_PROMPT = """You are a Saudi real estate document parser. Extract structured data from this land opportunity document.

Return ONLY a JSON object with these fields (use null for missing):
{
  "district": "الحي (Arabic name)",
  "city": "المدينة",
  "plan_number": "رقم المخطط (just the number)",
  "land_area_sqm": 88567.67,
  "land_status": "خام or مطور",
  "building_code": "نظام البناء code if mentioned",
  "boundaries": {
    "north": {"description": "text", "length_m": 84},
    "south": {"description": "text", "length_m": 351.34},
    "east": {"description": "text", "length_m": 189.23},
    "west": {"description": "text", "length_m": 312.83}
  },
  "google_maps_url": "full URL or short URL if present",
  "deed_number": "رقم الصك",
  "deed_reference": "المرجع",
  "survey_coordinates": [
    {"easting": 659155.96, "northing": 2757737.43},
    ...
  ],
  "property_type": "type description",
  "date": "date if present",
  "notes": "any additional notes"
}

Parse Arabic text carefully. Extract numbers from Arabic text including م² areas.
If survey coordinates are in a table, extract all rows as easting/northing pairs.
"""


# ---------------------------------------------------------------------------
# Step 1: Parse .docx
# ---------------------------------------------------------------------------

def parse_docx(file_bytes: bytes) -> str:
    """Extract all text from a .docx file."""
    doc = Document(io.BytesIO(file_bytes))

    parts: list[str] = []

    # Paragraphs
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    # Tables
    for table in doc.tables:
        parts.append("\n--- TABLE ---")
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            parts.append(" | ".join(cells))

    return "\n".join(parts)


# ---------------------------------------------------------------------------
# Step 2: Claude extraction
# ---------------------------------------------------------------------------

async def extract_fields(
    text: str,
    anthropic_client: AsyncAnthropic,
    model: str = "claude-sonnet-4-20250514",
) -> dict[str, Any]:
    """Use Claude to extract structured fields from document text."""
    response = await anthropic_client.messages.create(
        model=model,
        max_tokens=2048,
        system=EXTRACTION_PROMPT,
        messages=[{"role": "user", "content": f"Document text:\n\n{text[:8000]}"}],
    )

    # Extract JSON from response
    result_text = ""
    for block in response.content:
        if hasattr(block, "text"):
            result_text += block.text

    # Parse JSON (Claude sometimes wraps in ```json```)
    json_match = re.search(r'\{[\s\S]*\}', result_text)
    if json_match:
        try:
            return json.loads(json_match.group())
        except json.JSONDecodeError:
            log.warning("Failed to parse Claude JSON: %s", result_text[:300])

    return {"raw_response": result_text, "error": "Could not parse structured data"}


# ---------------------------------------------------------------------------
# Step 3: Resolve coordinates
# ---------------------------------------------------------------------------

async def resolve_coordinates(
    extracted: dict,
    http_client: httpx.AsyncClient,
) -> tuple[float, float] | None:
    """Resolve coordinates from Google Maps URL or survey points."""

    # Priority 1: Google Maps URL
    gmap_url = extracted.get("google_maps_url")
    if gmap_url:
        # Resolve short URLs
        if "goo.gl" in gmap_url or "maps.app" in gmap_url:
            try:
                r = await http_client.get(gmap_url, follow_redirects=True, timeout=10)
                gmap_url = str(r.url)
            except Exception as exc:
                log.warning("URL resolve failed: %s", exc)

        coords = parse_coordinates(gmap_url)
        if coords:
            log.info("Resolved from Google Maps: %.6f, %.6f", *coords)
            return coords

    # Priority 2: Survey coordinates (AIN-ABD UTM → WGS84)
    survey = extracted.get("survey_coordinates")
    if survey and len(survey) > 0:
        try:
            wgs_points = []
            for pt in survey:
                e = pt.get("easting")
                n = pt.get("northing")
                if e and n:
                    lng, lat = _utm_to_wgs84.transform(float(e), float(n))
                    wgs_points.append((lat, lng))

            if wgs_points:
                # Centroid
                avg_lat = sum(p[0] for p in wgs_points) / len(wgs_points)
                avg_lng = sum(p[1] for p in wgs_points) / len(wgs_points)
                log.info("Resolved from survey points (%d pts): %.6f, %.6f",
                         len(wgs_points), avg_lat, avg_lng)
                return (avg_lat, avg_lng)
        except Exception as exc:
            log.warning("Survey coordinate conversion failed: %s", exc)

    return None


# ---------------------------------------------------------------------------
# Step 4: Merge document + Geoportal
# ---------------------------------------------------------------------------

def merge_document_and_geoportal(
    extracted: dict,
    geoportal: dict | None,
) -> dict[str, Any]:
    """Merge extracted document data with Geoportal data, flagging conflicts.

    CRITICAL: Document area is authoritative for the DEAL.
    Geoportal identify often hits the parent plan boundary, not the subdivision plot.
    """
    doc_area = extracted.get("land_area_sqm")
    geo_area = geoportal.get("area_sqm") if geoportal else None

    # Document area is authoritative; Geoportal kept for reference
    authoritative_area = doc_area or geo_area
    area_source = "document" if doc_area else "geoportal"

    merged: dict[str, Any] = {
        "source": "document",
        "district": extracted.get("district"),
        "plan_number": extracted.get("plan_number"),
        "land_area_sqm": authoritative_area,
        "area_source": area_source,
        "document_area_sqm": doc_area,
        "land_status": extracted.get("land_status"),
        "building_code": extracted.get("building_code"),
        "boundaries": extracted.get("boundaries"),
        "deed_number": extracted.get("deed_number"),
        "deed_reference": extracted.get("deed_reference"),
        "property_type": extracted.get("property_type"),
        "notes": extracted.get("notes"),
    }

    conflicts: list[str] = []

    if geoportal:
        merged["geoportal_parcel_id"] = geoportal.get("parcel_id")
        merged["geoportal_district"] = geoportal.get("district_name")
        merged["geoportal_plan"] = geoportal.get("plan_number")
        merged["geoportal_area"] = geo_area
        merged["geoportal_building_code"] = geoportal.get("building_code_label")
        merged["geoportal_regulations"] = geoportal.get("regulations")
        merged["geoportal_market"] = geoportal.get("market")

        # Area mismatch warning
        if doc_area and geo_area:
            delta = abs(float(doc_area) - float(geo_area)) / float(doc_area)
            if delta > 0.05:
                conflicts.append(
                    f"\u26a0\ufe0f \u0627\u0644\u062c\u064a\u0648\u0628\u0648\u0631\u062a\u0627\u0644 \u064a\u0638\u0647\u0631 {geo_area:,.0f} \u0645\u00b2 "
                    f"(\u0627\u0644\u0645\u062e\u0637\u0637 \u0627\u0644\u0623\u0645) \u2014 "
                    f"\u0627\u0644\u0646\u0638\u0627\u0645 \u064a\u0633\u062a\u062e\u062f\u0645 {doc_area:,.0f} \u0645\u00b2 "
                    f"(\u0645\u0646 \u0645\u0633\u062a\u0646\u062f \u0627\u0644\u0639\u0631\u0636)"
                )

        doc_plan = str(extracted.get("plan_number", ""))
        geo_plan = str(geoportal.get("plan_number", ""))
        if doc_plan and geo_plan and doc_plan != geo_plan:
            conflicts.append(
                f"Plan mismatch: document says {doc_plan}, Geoportal says {geo_plan}."
            )

        geo_code = geoportal.get("building_code_label", "")
        if geo_code and "\u0645\u0631\u0627\u062c\u0639\u0629" in geo_code:
            conflicts.append(
                f"Building code: '{geo_code}' \u2014 raw land, no zoning assigned. "
                f"You must set FAR/floors manually."
            )

    merged["conflicts"] = conflicts
    return merged
